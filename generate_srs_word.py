from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_srs_document():
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Software Requirements Specification (SRS)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Features Implementation Summary', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # System Name
    system_name = doc.add_paragraph()
    system_name.add_run('Vanatvam Property Booking Management System').bold = True
    
    # Document Information
    doc.add_heading('Document Information', 1)
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    info_table.cell(0, 0).text = 'System Name:'
    info_table.cell(0, 1).text = 'Vanatvam Property Booking Management System'
    info_table.cell(1, 0).text = 'Document Type:'
    info_table.cell(1, 1).text = 'Features Implementation Summary'
    info_table.cell(2, 0).text = 'Version:'
    info_table.cell(2, 1).text = '1.0'
    info_table.cell(3, 0).text = 'Date:'
    info_table.cell(3, 1).text = 'Generated from Codebase Analysis'
    
    doc.add_paragraph()
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', 1)
    doc.add_paragraph(
        'The Vanatvam Property Booking Management System is a comprehensive web-based application '
        'designed for managing property bookings with a quota-based credit system. The system supports '
        'two primary user roles: Administrators and Property Owners, each with distinct functionalities and access levels.'
    )
    
    doc.add_paragraph('Technology Stack:', style='List Bullet')
    doc.add_paragraph('Frontend: React with TypeScript', style='List Bullet 2')
    doc.add_paragraph('Backend: Python FastAPI', style='List Bullet 2')
    doc.add_paragraph('Database: PostgreSQL', style='List Bullet 2')
    doc.add_paragraph('Authentication: JWT-based with bcrypt password hashing', style='List Bullet 2')
    
    # 2. Authentication & Authorization
    doc.add_heading('2. Authentication & Authorization Features', 1)
    
    doc.add_heading('2.1 User Registration', 2)
    doc.add_paragraph('• Self-registration for property owners')
    doc.add_paragraph('• Email verification system with token-based verification')
    doc.add_paragraph('• Email verification link sent upon registration')
    doc.add_paragraph('• 24-hour token expiration')
    doc.add_paragraph('• Status: PENDING until admin approval')
    
    doc.add_heading('2.2 User Login', 2)
    doc.add_paragraph('• Email and password authentication')
    doc.add_paragraph('• JWT token generation (30-minute expiration)')
    doc.add_paragraph('• Email verification required before login (for owners)')
    doc.add_paragraph('• Admin users bypass email verification')
    doc.add_paragraph('• Status-based access control (only ACTIVE users can login)')
    
    doc.add_heading('2.3 Password Management', 2)
    doc.add_paragraph('• Password reset functionality with token-based reset links')
    doc.add_paragraph('• 1-hour token expiration for password reset')
    doc.add_paragraph('• Secure password hashing using bcrypt')
    doc.add_paragraph('• Forgot password flow with email notification')
    
    doc.add_heading('2.4 Email Verification', 2)
    doc.add_paragraph('• Email verification endpoint with token validation')
    doc.add_paragraph('• Verification notification email')
    doc.add_paragraph('• Email verification status tracking')
    doc.add_paragraph('• Token expiration handling')
    
    doc.add_heading('2.5 Role-Based Access Control', 2)
    doc.add_paragraph('• Two user roles: ADMIN and OWNER')
    doc.add_paragraph('• Three user statuses: PENDING, ACTIVE, SUSPENDED')
    doc.add_paragraph('• Property-based data isolation for owners')
    doc.add_paragraph('• Admin-only endpoints protection')
    
    # 3. Admin Module
    doc.add_heading('3. Admin Module Features', 1)
    
    doc.add_heading('3.1 Community Management (Member Management)', 2)
    
    doc.add_heading('ADM-01: Pending Member Queue', 3)
    doc.add_paragraph('• View all pending member registrations')
    doc.add_paragraph('• Filter by email verification status')
    doc.add_paragraph('• Display member registration details')
    doc.add_paragraph('• Queue management interface')
    
    doc.add_heading('ADM-02: Member Activation & Assignment', 3)
    doc.add_paragraph('• Activate pending members')
    doc.add_paragraph('• Assign members to properties')
    doc.add_paragraph('• Set initial quota (weekday/weekend)')
    doc.add_paragraph('• Default quotas: 12 weekday, 6 weekend credits')
    doc.add_paragraph('• Automatic quota balance initialization')
    doc.add_paragraph('• Transaction record creation for activation')
    doc.add_paragraph('• Email notification on activation')
    
    doc.add_heading('ADM-03: Member Lookup & History', 3)
    doc.add_paragraph('• Search members by name, email, phone, or property')
    doc.add_paragraph('• View member details (profile, bookings, transactions)')
    doc.add_paragraph('• Complete member history tracking')
    doc.add_paragraph('• Get all active members list')
    doc.add_paragraph('• Edit member credentials (name, email, phone, password)')
    
    doc.add_heading('ADM-04: Manual Quota Adjustment', 3)
    doc.add_paragraph('• Adjust quota balances manually')
    doc.add_paragraph('• Add or subtract weekday/weekend credits')
    doc.add_paragraph('• Add adjustment notes/descriptions')
    doc.add_paragraph('• Transaction history recording')
    doc.add_paragraph('• View all quota adjustment history')
    doc.add_paragraph('• Prevent negative balances')
    
    doc.add_heading('ADM-05: Member Deactivation', 3)
    doc.add_paragraph('• Suspend member accounts')
    doc.add_paragraph('• Reactivate suspended members')
    doc.add_paragraph('• Delete member accounts (with cascading deletion)')
    doc.add_paragraph('• Automatic cleanup of related records')
    
    doc.add_heading('Member Rejection', 3)
    doc.add_paragraph('• Reject pending member registrations')
    doc.add_paragraph('• Send rejection email notifications')
    doc.add_paragraph('• Delete rejected user records')
    
    doc.add_heading('3.2 Inventory Governance (Property & Cottage Management)', 2)
    
    doc.add_heading('ADM-06: Property Management', 3)
    doc.add_paragraph('• Create properties/sanctuaries (e.g., Madhuvana)')
    doc.add_paragraph('• List all properties')
    doc.add_paragraph('• Update property details (name, description)')
    doc.add_paragraph('• Property deletion support')
    
    doc.add_heading('ADM-07: Cottage Inventory Management', 3)
    doc.add_paragraph('• Create cottages under properties')
    doc.add_paragraph('• Set cottage ID (e.g., "C-12")')
    doc.add_paragraph('• Define cottage capacity')
    doc.add_paragraph('• Add amenities information')
    doc.add_paragraph('• List cottages (filterable by property)')
    doc.add_paragraph('• Update cottage details')
    doc.add_paragraph('• Delete cottage records')
    
    doc.add_heading('ADM-08: Maintenance Blocking', 3)
    doc.add_paragraph('• Create maintenance blocks for cottages')
    doc.add_paragraph('• Set maintenance date ranges')
    doc.add_paragraph('• Add maintenance reasons/notes')
    doc.add_paragraph('• Auto-reject pending bookings during maintenance')
    doc.add_paragraph('• Automatic credit refunds for rejected bookings')
    doc.add_paragraph('• View all maintenance blocks')
    doc.add_paragraph('• Update maintenance blocks')
    doc.add_paragraph('• Delete maintenance blocks')
    doc.add_paragraph('• View bookings affected by maintenance blocks')
    doc.add_paragraph('• Bulk revoke bookings for maintenance periods')
    
    doc.add_heading('ADM-09: Inventory Health View', 3)
    doc.add_paragraph('• Calendar view of cottage availability')
    doc.add_paragraph('• Filter by property and date range')
    doc.add_paragraph('• Status indicators: Available, Booked, Maintenance')
    doc.add_paragraph('• Comprehensive inventory status tracking')
    
    doc.add_heading('3.3 Booking Governance', 2)
    
    doc.add_heading('ADM-10: Approval Queue', 3)
    doc.add_paragraph('• View all pending booking requests')
    doc.add_paragraph('• Display booking details (user, cottage, dates, credits)')
    doc.add_paragraph('• Sort by creation date')
    doc.add_paragraph('• Property and cottage information display')
    
    doc.add_heading('ADM-11: Booking Decision Workflow', 3)
    doc.add_paragraph('• Approve booking requests')
    doc.add_paragraph('• Reject booking requests with notes')
    doc.add_paragraph('• Automatic credit handling (approved = credits deducted, rejected = refunded)')
    doc.add_paragraph('• Transaction history for decisions')
    doc.add_paragraph('• Decision notes storage')
    
    doc.add_heading('ADM-12: Emergency Revocation', 3)
    doc.add_paragraph('• Revoke confirmed bookings')
    doc.add_paragraph('• Add revocation reason')
    doc.add_paragraph('• Full credit refund on revocation')
    doc.add_paragraph('• Transaction history recording')
    doc.add_paragraph('• Bulk revoke bookings for maintenance blocks')
    
    doc.add_heading('ADM-13: Admin Override Booking', 3)
    doc.add_paragraph('• Create bookings directly (bypass standard flow)')
    doc.add_paragraph('• Override availability checks')
    doc.add_paragraph('• Direct confirmation (no approval needed)')
    
    doc.add_heading('Booking Calendar View', 3)
    doc.add_paragraph('• View all bookings (confirmed and pending)')
    doc.add_paragraph('• Calendar display with cottage names')
    doc.add_paragraph('• User and property information')
    doc.add_paragraph('• Status indicators')
    
    doc.add_heading('Rejected Bookings View', 3)
    doc.add_paragraph('• View all rejected and cancelled bookings')
    doc.add_paragraph('• Display owner, sanctuary, and cottage details')
    doc.add_paragraph('• Decision notes and timestamps')
    doc.add_paragraph('• Chronological sorting')
    
    doc.add_heading('3.4 Holiday & Season Configuration', 2)
    
    doc.add_heading('ADM-14: Holiday Configuration', 3)
    doc.add_paragraph('• Set holiday dates')
    doc.add_paragraph('• Add holiday names')
    doc.add_paragraph('• View all configured holidays')
    doc.add_paragraph('• Update holiday details')
    doc.add_paragraph('• Delete holidays')
    doc.add_paragraph('• Holidays use weekend pricing automatically')
    
    doc.add_heading('ADM-15: Peak Season Definition', 3)
    doc.add_paragraph('• Create peak season periods')
    doc.add_paragraph('• Define start and end dates')
    doc.add_paragraph('• Name peak seasons')
    doc.add_paragraph('• View all peak seasons')
    doc.add_paragraph('• Update peak season definitions')
    doc.add_paragraph('• Delete peak seasons')
    doc.add_paragraph('• Peak season dates use weekend pricing')
    
    doc.add_heading('3.5 Quota Management', 2)
    
    doc.add_heading('ADM-16: Global Quota Reset', 3)
    doc.add_paragraph('• Reset all active user quotas')
    doc.add_paragraph('• Annual quota reset functionality')
    doc.add_paragraph('• Set balances to quota amounts')
    doc.add_paragraph('• Transaction history for resets')
    doc.add_paragraph('• Batch processing of all users')
    
    doc.add_heading('3.6 Email Configuration', 2)
    
    doc.add_heading('Email Service Configuration', 3)
    doc.add_paragraph('• Configure SMTP server settings')
    doc.add_paragraph('• Set SMTP port, username, password')
    doc.add_paragraph('• Configure sender email address')
    doc.add_paragraph('• Set frontend URL for email links')
    doc.add_paragraph('• Enable/disable email service')
    doc.add_paragraph('• Test email configuration')
    doc.add_paragraph('• Secure password storage')
    
    doc.add_heading('Email Template Management', 3)
    doc.add_paragraph('• Create email templates')
    doc.add_paragraph('• Template types: registration, verification, approval, rejection')
    doc.add_paragraph('• HTML and plain text template support')
    doc.add_paragraph('• Update email templates')
    doc.add_paragraph('• View all templates')
    doc.add_paragraph('• Customizable email content')
    
    doc.add_heading('3.7 Reporting & Analytics', 2)
    
    doc.add_heading('Reports & Statistics', 3)
    doc.add_paragraph('• User statistics (total, active, pending, suspended)')
    doc.add_paragraph('• Booking statistics (total, confirmed, pending, rejected, cancelled)')
    doc.add_paragraph('• User registration trends (last 12 months)')
    doc.add_paragraph('• Booking creation trends (last 12 months)')
    doc.add_paragraph('• Pie charts data for bookings by status')
    doc.add_paragraph('• Pie charts data for users by status')
    doc.add_paragraph('• Time-series data for analytics')
    
    doc.add_heading('Audit Trail', 3)
    doc.add_paragraph('• Comprehensive system activity log')
    doc.add_paragraph('• Transaction history tracking')
    doc.add_paragraph('• Booking status change tracking')
    doc.add_paragraph('• Member activation tracking')
    doc.add_paragraph('• Filterable by date range')
    doc.add_paragraph('• Detailed activity descriptions')
    doc.add_paragraph('• User and property information')
    doc.add_paragraph('• Chronological sorting')
    
    doc.add_heading('3.8 Admin User Management', 2)
    
    doc.add_heading('Create Admin Users', 3)
    doc.add_paragraph('• Create additional admin accounts')
    doc.add_paragraph('• Only existing admins can create admins')
    doc.add_paragraph('• Full admin privileges assignment')
    doc.add_paragraph('• Email uniqueness validation')
    
    # 4. Owner Module
    doc.add_heading('4. Owner Module Features', 1)
    
    doc.add_heading('4.1 Access & Identity', 2)
    
    doc.add_heading('OWN-01: Self-Registration', 3)
    doc.add_paragraph('• User registration form')
    doc.add_paragraph('• Email, phone, name, password input')
    doc.add_paragraph('• Email verification requirement')
    doc.add_paragraph('• Pending status until admin approval')
    
    doc.add_heading('OWN-02: Account Status Awareness', 3)
    doc.add_paragraph('• Pending status display')
    doc.add_paragraph('• Account activation notification')
    doc.add_paragraph('• Access restrictions for pending users')
    
    doc.add_heading('OWN-03: Property Context', 3)
    doc.add_paragraph('• Property information display')
    doc.add_paragraph('• Cottage listing for assigned property')
    doc.add_paragraph('• Property-based dashboard')
    doc.add_paragraph('• Access limited to assigned property only')
    
    doc.add_heading('4.2 Booking Calendar & Availability', 2)
    
    doc.add_heading('OWN-04: Real-Time Availability', 3)
    doc.add_paragraph('• Cottage availability calendar view')
    doc.add_paragraph('• Date range selection')
    doc.add_paragraph('• Availability status indicators: Available dates, Booked dates, Maintenance blocked dates, Holiday indicators, Peak season indicators')
    doc.add_paragraph('• Color-coded calendar display')
    
    doc.add_heading('OWN-05: Holiday & Peak Season Transparency', 3)
    doc.add_paragraph('• Visual indicators for holidays')
    doc.add_paragraph('• Peak season date highlighting')
    doc.add_paragraph('• Pricing transparency (weekend pricing for holidays/peak seasons)')
    doc.add_paragraph('• Clear date labeling')
    
    doc.add_heading('OWN-06: Cost Calculator', 3)
    doc.add_paragraph('• Pre-booking cost calculation')
    doc.add_paragraph('• Weekday vs weekend credit calculation')
    doc.add_paragraph('• Holiday pricing calculation')
    doc.add_paragraph('• Peak season pricing calculation')
    doc.add_paragraph('• Detailed cost breakdown')
    doc.add_paragraph('• Day-by-day credit calculation')
    
    doc.add_heading('OWN-07: Submit Booking Request', 3)
    doc.add_paragraph('• Create booking requests')
    doc.add_paragraph('• Date range selection')
    doc.add_paragraph('• Cottage selection')
    doc.add_paragraph('• Automatic availability validation')
    doc.add_paragraph('• Credit escrow system (credits deducted on submission)')
    doc.add_paragraph('• Insufficient credit validation')
    doc.add_paragraph('• Booking status: PENDING')
    doc.add_paragraph('• Transaction record creation')
    doc.add_paragraph('• Conflict checking (bookings, maintenance)')
    
    doc.add_heading('4.3 Quota Management', 2)
    
    doc.add_heading('OWN-08: Balance Dashboard', 3)
    doc.add_paragraph('• Weekday quota and balance display')
    doc.add_paragraph('• Weekend quota and balance display')
    doc.add_paragraph('• Available credits calculation')
    doc.add_paragraph('• Pending booking credits (escrow) display')
    doc.add_paragraph('• Confirmed booking credits display')
    doc.add_paragraph('• Visual quota status indicators')
    
    doc.add_heading('OWN-09: Escrow Visibility', 3)
    doc.add_paragraph('• Pending booking credits tracking')
    doc.add_paragraph('• Escrowed credits display')
    doc.add_paragraph('• Available vs escrowed credit separation')
    doc.add_paragraph('• Real-time balance updates')
    
    doc.add_heading('OWN-10: Transaction History', 3)
    doc.add_paragraph('• Complete transaction ledger')
    doc.add_paragraph('• Transaction types: booking, refund, reset, manual_adjustment, activation')
    doc.add_paragraph('• Credit change tracking (weekday/weekend)')
    doc.add_paragraph('• Transaction descriptions')
    doc.add_paragraph('• Booking association')
    doc.add_paragraph('• Chronological sorting')
    doc.add_paragraph('• Timestamp information')
    
    doc.add_heading('4.4 Trip Management (My Stays)', 2)
    
    doc.add_heading('OWN-11: Trip Status Tracking', 3)
    doc.add_paragraph('• View all bookings (all statuses)')
    doc.add_paragraph('• Filter by booking status')
    doc.add_paragraph('• Booking details display: Cottage information, Check-in/check-out dates, Credits used, Booking status, Creation timestamp')
    doc.add_paragraph('• Chronological sorting')
    
    doc.add_heading('OWN-12: Self-Cancellation', 3)
    doc.add_paragraph('• Cancel own bookings')
    doc.add_paragraph('• Automatic credit refund')
    doc.add_paragraph('• Transaction history recording')
    doc.add_paragraph('• Status validation (cannot cancel already cancelled/rejected)')
    doc.add_paragraph('• Full credit restoration')
    
    doc.add_heading('OWN-13: Booking Receipt', 3)
    doc.add_paragraph('• Detailed booking receipt generation')
    doc.add_paragraph('• Property and cottage information')
    doc.add_paragraph('• Check-in/check-out dates')
    doc.add_paragraph('• Duration calculation')
    doc.add_paragraph('• Credits breakdown (weekday/weekend)')
    doc.add_paragraph('• Total credits used')
    doc.add_paragraph('• Booking status')
    doc.add_paragraph('• Decision notes (if applicable)')
    doc.add_paragraph('• Guest rules and information')
    
    # 5. System Features
    doc.add_heading('5. System Features', 1)
    
    doc.add_heading('5.1 Quota System', 2)
    doc.add_paragraph('• Annual Quota Allocation: Default 12 weekday / 6 weekend credits per user')
    doc.add_paragraph('• Escrow System: Credits deducted when booking is submitted (pending status)')
    doc.add_paragraph('• Automatic Refunds: Credits refunded on rejection or cancellation')
    doc.add_paragraph('• Transaction Tracking: All quota changes recorded in transaction history')
    doc.add_paragraph('• Balance Management: Real-time balance updates')
    doc.add_paragraph('• Quota Types: Separate weekday and weekend quotas')
    doc.add_paragraph('• Negative Balance Prevention: System prevents negative balances')
    
    doc.add_heading('5.2 Booking Workflow', 2)
    p = doc.add_paragraph()
    p.add_run('1. ').bold = True
    p.add_run('Owner selects dates and cottage')
    p = doc.add_paragraph()
    p.add_run('2. ').bold = True
    p.add_run('System calculates cost (weekday vs weekend credits)')
    p = doc.add_paragraph()
    p.add_run('3. ').bold = True
    p.add_run('Credits are escrowed (deducted from balance)')
    p = doc.add_paragraph()
    p.add_run('4. ').bold = True
    p.add_run('Booking status: PENDING')
    p = doc.add_paragraph()
    p.add_run('5. ').bold = True
    p.add_run('Admin reviews in approval queue')
    p = doc.add_paragraph()
    p.add_run('6. ').bold = True
    p.add_run('Admin approves → CONFIRMED (credits remain deducted)')
    p = doc.add_paragraph()
    p.add_run('7. ').bold = True
    p.add_run('Admin rejects → REJECTED (credits refunded)')
    
    doc.add_heading('5.3 Pricing Logic', 2)
    doc.add_paragraph('• Weekday Pricing: Standard weekdays (Monday-Friday, non-holiday, non-peak)')
    doc.add_paragraph('• Weekend Pricing: Weekends (Saturday-Sunday), holidays, and peak season dates')
    doc.add_paragraph('• Holiday Pricing: Holidays automatically use weekend credit pricing')
    doc.add_paragraph('• Peak Season Pricing: Peak season dates use weekend credit pricing')
    doc.add_paragraph('• Transparent Calculation: Users see breakdown of weekday/weekend/holiday/peak days')
    
    doc.add_heading('5.4 Maintenance Blocking', 2)
    doc.add_paragraph('• Blocks specific date ranges for cottages')
    doc.add_paragraph('• Auto-rejects pending bookings in blocked dates')
    doc.add_paragraph('• Prevents new bookings during maintenance')
    doc.add_paragraph('• Automatic credit refunds')
    doc.add_paragraph('• Maintenance reason tracking')
    doc.add_paragraph('• Bulk booking revocation support')
    
    doc.add_heading('5.5 Email Notification System', 2)
    doc.add_paragraph('• Registration confirmation emails')
    doc.add_paragraph('• Email verification links')
    doc.add_paragraph('• Account activation notifications')
    doc.add_paragraph('• Booking approval notifications')
    doc.add_paragraph('• Booking rejection notifications')
    doc.add_paragraph('• Member rejection notifications')
    doc.add_paragraph('• Configurable SMTP settings')
    doc.add_paragraph('• Customizable email templates')
    doc.add_paragraph('• Test email functionality')
    
    # 6. Security Features
    doc.add_heading('6. Security Features', 1)
    
    doc.add_heading('6.1 Authentication Security', 2)
    doc.add_paragraph('• JWT-based authentication')
    doc.add_paragraph('• Password hashing with bcrypt')
    doc.add_paragraph('• Token expiration (30 minutes for access tokens)')
    doc.add_paragraph('• Secure token generation')
    
    doc.add_heading('6.2 Authorization Security', 2)
    doc.add_paragraph('• Role-based access control (admin/owner)')
    doc.add_paragraph('• Status-based access (pending/active/suspended)')
    doc.add_paragraph('• Property-based data isolation')
    doc.add_paragraph('• Endpoint-level authorization checks')
    doc.add_paragraph('• Admin-only endpoint protection')
    
    doc.add_heading('6.3 Data Security', 2)
    doc.add_paragraph('• Password encryption')
    doc.add_paragraph('• Email verification tokens (24-hour expiration)')
    doc.add_paragraph('• Password reset tokens (1-hour expiration)')
    doc.add_paragraph('• Secure token storage')
    doc.add_paragraph('• SQL injection prevention (SQLAlchemy ORM)')
    doc.add_paragraph('• Input validation (Pydantic schemas)')
    
    doc.add_heading('6.4 API Security', 2)
    doc.add_paragraph('• CORS configuration')
    doc.add_paragraph('• Credential-based authentication')
    doc.add_paragraph('• Protected API endpoints')
    doc.add_paragraph('• Session management')
    
    # 7. Database Features
    doc.add_heading('7. Database Features', 1)
    
    doc.add_heading('7.1 Data Models', 2)
    doc.add_paragraph('• User Model: Users with roles, status, quota, property assignment')
    doc.add_paragraph('• Property Model: Properties/sanctuaries')
    doc.add_paragraph('• Cottage Model: Individual cottages under properties')
    doc.add_paragraph('• Booking Model: Booking requests with status tracking')
    doc.add_paragraph('• MaintenanceBlock Model: Maintenance date ranges')
    doc.add_paragraph('• SystemCalendar Model: Holiday and peak season dates')
    doc.add_paragraph('• PeakSeason Model: Peak season definitions')
    doc.add_paragraph('• QuotaTransaction Model: Transaction history')
    doc.add_paragraph('• EmailConfig Model: Email service configuration')
    doc.add_paragraph('• EmailTemplate Model: Email template storage')
    
    doc.add_heading('7.2 Data Relationships', 2)
    doc.add_paragraph('• Users belong to Properties (many-to-one)')
    doc.add_paragraph('• Cottages belong to Properties (many-to-one)')
    doc.add_paragraph('• Bookings belong to Users and Cottages (many-to-one each)')
    doc.add_paragraph('• Maintenance blocks belong to Cottages (many-to-one)')
    doc.add_paragraph('• Transactions belong to Users (many-to-one)')
    doc.add_paragraph('• Foreign key relationships for data integrity')
    
    doc.add_heading('7.3 Data Integrity', 2)
    doc.add_paragraph('• Unique constraints (email, property names)')
    doc.add_paragraph('• Foreign key constraints')
    doc.add_paragraph('• Enum types for status fields')
    doc.add_paragraph('• Timestamp tracking (created_at, updated_at)')
    doc.add_paragraph('• Cascading deletions (where appropriate)')
    
    # 8. Implementation Status
    doc.add_heading('8. Implementation Status', 1)
    doc.add_paragraph(
        'All features listed in this document have been FULLY IMPLEMENTED in the codebase, including:'
    )
    doc.add_paragraph('✓ Backend API endpoints', style='List Bullet')
    doc.add_paragraph('✓ Frontend React components', style='List Bullet')
    doc.add_paragraph('✓ Database models and relationships', style='List Bullet')
    doc.add_paragraph('✓ Authentication and authorization', style='List Bullet')
    doc.add_paragraph('✓ Email notification system', style='List Bullet')
    doc.add_paragraph('✓ Quota management system', style='List Bullet')
    doc.add_paragraph('✓ Booking workflow', style='List Bullet')
    doc.add_paragraph('✓ Admin dashboard', style='List Bullet')
    doc.add_paragraph('✓ Owner dashboard', style='List Bullet')
    doc.add_paragraph('✓ Calendar interfaces', style='List Bullet')
    doc.add_paragraph('✓ Reporting and analytics', style='List Bullet')
    doc.add_paragraph('✓ Audit trail', style='List Bullet')
    
    # 9. System Architecture
    doc.add_heading('9. System Architecture', 1)
    doc.add_paragraph('• Frontend: React application with TypeScript')
    doc.add_paragraph('• Backend: FastAPI REST API')
    doc.add_paragraph('• Database: PostgreSQL with SQLAlchemy ORM')
    doc.add_paragraph('• Authentication: JWT tokens')
    doc.add_paragraph('• Email: SMTP-based email service')
    doc.add_paragraph('• Deployment: Docker support (docker-compose.yml)')
    
    # Appendix
    doc.add_page_break()
    doc.add_heading('Appendix A: Feature Identifiers', 1)
    
    doc.add_heading('Admin Features', 2)
    features_table = doc.add_table(rows=17, cols=2)
    features_table.style = 'Light Grid Accent 1'
    features_table.cell(0, 0).text = 'ADM-01'
    features_table.cell(0, 1).text = 'Pending Member Queue'
    features_table.cell(1, 0).text = 'ADM-02'
    features_table.cell(1, 1).text = 'Member Activation & Assignment'
    features_table.cell(2, 0).text = 'ADM-03'
    features_table.cell(2, 1).text = 'Member Lookup & History'
    features_table.cell(3, 0).text = 'ADM-04'
    features_table.cell(3, 1).text = 'Manual Quota Adjustment'
    features_table.cell(4, 0).text = 'ADM-05'
    features_table.cell(4, 1).text = 'Member Deactivation'
    features_table.cell(5, 0).text = 'ADM-06'
    features_table.cell(5, 1).text = 'Property Management'
    features_table.cell(6, 0).text = 'ADM-07'
    features_table.cell(6, 1).text = 'Cottage Inventory Management'
    features_table.cell(7, 0).text = 'ADM-08'
    features_table.cell(7, 1).text = 'Maintenance Blocking'
    features_table.cell(8, 0).text = 'ADM-09'
    features_table.cell(8, 1).text = 'Inventory Health View'
    features_table.cell(9, 0).text = 'ADM-10'
    features_table.cell(9, 1).text = 'Approval Queue'
    features_table.cell(10, 0).text = 'ADM-11'
    features_table.cell(10, 1).text = 'Booking Decision Workflow'
    features_table.cell(11, 0).text = 'ADM-12'
    features_table.cell(11, 1).text = 'Emergency Revocation'
    features_table.cell(12, 0).text = 'ADM-13'
    features_table.cell(12, 1).text = 'Admin Override Booking'
    features_table.cell(13, 0).text = 'ADM-14'
    features_table.cell(13, 1).text = 'Holiday Configuration'
    features_table.cell(14, 0).text = 'ADM-15'
    features_table.cell(14, 1).text = 'Peak Season Definition'
    features_table.cell(15, 0).text = 'ADM-16'
    features_table.cell(15, 1).text = 'Global Quota Reset'
    
    doc.add_paragraph()
    doc.add_heading('Owner Features', 2)
    owner_table = doc.add_table(rows=14, cols=2)
    owner_table.style = 'Light Grid Accent 1'
    owner_table.cell(0, 0).text = 'OWN-01'
    owner_table.cell(0, 1).text = 'Self-Registration'
    owner_table.cell(1, 0).text = 'OWN-02'
    owner_table.cell(1, 1).text = 'Account Status Awareness'
    owner_table.cell(2, 0).text = 'OWN-03'
    owner_table.cell(2, 1).text = 'Property Context'
    owner_table.cell(3, 0).text = 'OWN-04'
    owner_table.cell(3, 1).text = 'Real-Time Availability'
    owner_table.cell(4, 0).text = 'OWN-05'
    owner_table.cell(4, 1).text = 'Holiday & Peak Season Transparency'
    owner_table.cell(5, 0).text = 'OWN-06'
    owner_table.cell(5, 1).text = 'Cost Calculator'
    owner_table.cell(6, 0).text = 'OWN-07'
    owner_table.cell(6, 1).text = 'Submit Booking Request'
    owner_table.cell(7, 0).text = 'OWN-08'
    owner_table.cell(7, 1).text = 'Balance Dashboard'
    owner_table.cell(8, 0).text = 'OWN-09'
    owner_table.cell(8, 1).text = 'Escrow Visibility'
    owner_table.cell(9, 0).text = 'OWN-10'
    owner_table.cell(9, 1).text = 'Transaction History'
    owner_table.cell(10, 0).text = 'OWN-11'
    owner_table.cell(10, 1).text = 'Trip Status Tracking'
    owner_table.cell(11, 0).text = 'OWN-12'
    owner_table.cell(11, 1).text = 'Self-Cancellation'
    owner_table.cell(12, 0).text = 'OWN-13'
    owner_table.cell(12, 1).text = 'Booking Receipt'
    
    # Save document
    doc.save('SRS_Vanatvam_Features.docx')
    print("SRS document created successfully: SRS_Vanatvam_Features.docx")

if __name__ == '__main__':
    create_srs_document()

